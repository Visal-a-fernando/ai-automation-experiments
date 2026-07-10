#!/usr/bin/env python3
"""
Incrementally write a troubleshooting case into a Word doc (.docx) that lives
in Visal's OneDrive sync folder. The wizard calls this once per phase as the
chat progresses, so the doc grows live and Visal can watch it from any
device with OneDrive access.

Sub-commands:

  start       Create the doc with the header and empty section skeleton.
  add-qa      Append a Q/A pair to the Conversation Log section.
  set-causes  Replace the Possible Causes section with a ranked list.
  add-step    Append one diagnostic step block (step + how-to + result +
              conclusion + outcome).
  set-flow    Replace the "How I Worked Through It" narrative.
  finalize    Fill the Root Cause / Resolution / Lesson Learned sections.

The doc has fixed section headings; each command finds its heading, deletes
the block that belongs to it, and writes fresh content. So calling a "set-"
command twice replaces, calling "add-" twice appends a second entry.

Style rules carried through: every string is run through clean() which
strips em dashes, en dashes, and double-hyphens. Visal's house rule.

Requires: python-docx (pip install python-docx)
"""
from __future__ import annotations

import argparse
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SECTION_ORDER = [
    "Symptom",
    "Environment",
    "Conversation Log",
    "Possible Causes",
    "Diagnostic Steps",
    "How I Worked Through It",
    "Root Cause",
    "Resolution",
    "Lesson Learned",
]

HEADING_STYLE = "Heading 1"
SUBHEADING_STYLE = "Heading 2"


def clean(text: str) -> str:
    if not text:
        return ""
    return (
        text.replace("—", ", ")  # em dash
        .replace("–", " to ")     # en dash
        .replace("--", ", ")
    )


def open_doc(path: Path) -> Document:
    if not path.exists():
        sys.exit(f"doc not found: {path}. Run 'start' first.")
    return Document(str(path))


def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(path))
    except PermissionError:
        sys.exit(
            f"could not save {path}. Is the doc open in Word on this laptop? "
            "Close it there (view from another device instead) and retry."
        )


def find_heading_index(doc: Document, heading_text: str) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == HEADING_STYLE and p.text.strip().lower() == heading_text.lower():
            return i
    return None


def next_heading_index(doc: Document, after: int) -> int:
    """Index of the next Heading 1 paragraph after `after`, or len(paragraphs)."""
    for i in range(after + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name == HEADING_STYLE:
            return i
    return len(doc.paragraphs)


def delete_section_body(doc: Document, heading_text: str) -> int:
    """Delete every paragraph between the heading and the next heading.
    Returns the index AFTER the heading where new content should be inserted."""
    h_idx = find_heading_index(doc, heading_text)
    if h_idx is None:
        return -1
    end = next_heading_index(doc, h_idx)
    # Delete in reverse so indexes stay valid.
    for i in range(end - 1, h_idx, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
    return h_idx + 1


def insert_paragraph_after(
    doc: Document,
    after_idx: int,
    text: str = "",
    style: str | None = None,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    size: int | None = None,
    left_indent: float | None = None,
):
    """Insert a paragraph immediately after paragraphs[after_idx]."""
    anchor = doc.paragraphs[after_idx]
    new_p = anchor._element.makeelement(qn("w:p"), {})
    anchor._element.addnext(new_p)
    # Re-grab via the doc so styles work.
    target_paragraph = None
    for p in doc.paragraphs:
        if p._element is new_p:
            target_paragraph = p
            break
    assert target_paragraph is not None
    if style:
        target_paragraph.style = doc.styles[style]
    if text:
        run = target_paragraph.add_run(clean(text))
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
    if left_indent is not None:
        target_paragraph.paragraph_format.left_indent = Pt(left_indent)
    return target_paragraph


def append_to_section(doc: Document, heading_text: str, render_fn) -> None:
    """Append content to the end of an existing section (before the next heading)."""
    h_idx = find_heading_index(doc, heading_text)
    if h_idx is None:
        sys.exit(f"section '{heading_text}' not found. Doc was not started properly.")
    end = next_heading_index(doc, h_idx)
    insert_after = end - 1  # last paragraph in this section (or the heading itself)
    render_fn(insert_after)


# ----------------------- sub-commands -----------------------


def cmd_start(args, doc_path: Path) -> None:
    if doc_path.exists():
        # Don't clobber an existing case doc.
        sys.exit(f"doc already exists at {doc_path}. Use a different path or delete it first.")
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = title.add_run("IT Support Case Record")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Meta line
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"{args.case_id}  |  {datetime.now().strftime('%Y-%m-%d')}  |  "
        f"Engineer: {clean(args.engineer)}  |  Client: {clean(args.client)}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Section skeleton
    for section in SECTION_ORDER:
        h = doc.add_paragraph(section, style=HEADING_STYLE)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        # Seed content for the two sections we already know
        if section == "Symptom" and args.symptom:
            doc.add_paragraph(clean(args.symptom))
        elif section == "Environment" and args.environment:
            doc.add_paragraph(clean(args.environment))
        else:
            doc.add_paragraph()  # empty placeholder so the heading isn't naked

    save_doc(doc, doc_path)
    print(f"[OK] Started case doc at {doc_path}")


def cmd_add_qa(args, doc_path: Path) -> None:
    doc = open_doc(doc_path)

    def render(after_idx: int):
        nonlocal_idx = [after_idx]

        def insert(text, **kw):
            p = insert_paragraph_after(doc, nonlocal_idx[0], text, **kw)
            # Update anchor index to the new paragraph's index
            for i, pp in enumerate(doc.paragraphs):
                if pp._element is p._element:
                    nonlocal_idx[0] = i
                    break

        insert(
            f"Q. {args.question}",
            bold=True,
            color=RGBColor(0x1F, 0x3A, 0x5F),
            size=10,
        )
        insert(
            f"A. {args.answer or '(not answered)'}",
            size=10,
            left_indent=18,
        )

    append_to_section(doc, "Conversation Log", render)
    save_doc(doc, doc_path)
    print("[OK] Added Q/A")


def cmd_set_causes(args, doc_path: Path) -> None:
    doc = open_doc(doc_path)
    insert_at = delete_section_body(doc, "Possible Causes")
    if insert_at < 0:
        sys.exit("Possible Causes heading missing")

    # cause/reason pair by order
    causes = args.cause or []
    reasons = (args.reason or []) + [""] * max(0, len(causes) - len(args.reason or []))

    cursor = insert_at - 1  # heading index
    for idx, (cause, reason) in enumerate(zip(causes, reasons), start=1):
        # Cause line
        p1 = insert_paragraph_after(
            doc, cursor, f"{idx}. {cause}", bold=True, size=10,
        )
        for i, pp in enumerate(doc.paragraphs):
            if pp._element is p1._element:
                cursor = i
                break
        # Reasoning line, indented
        if reason:
            p2 = insert_paragraph_after(
                doc, cursor, f"why: {reason}", italic=True, size=10, left_indent=18,
                color=RGBColor(0x3A, 0x3A, 0x3A),
            )
            for i, pp in enumerate(doc.paragraphs):
                if pp._element is p2._element:
                    cursor = i
                    break

    save_doc(doc, doc_path)
    print(f"[OK] Set {len(causes)} possible causes")


def cmd_add_step(args, doc_path: Path) -> None:
    doc = open_doc(doc_path)

    # Count existing steps in this section to derive the next step number.
    h_idx = find_heading_index(doc, "Diagnostic Steps")
    end = next_heading_index(doc, h_idx)
    step_count = sum(
        1 for i in range(h_idx + 1, end)
        if doc.paragraphs[i].style.name == SUBHEADING_STYLE
    )
    step_num = step_count + 1

    def render(after_idx: int):
        cursor = [after_idx]

        def insert(text, **kw):
            style = kw.pop("style", None)
            p = insert_paragraph_after(doc, cursor[0], text, style=style, **kw)
            for i, pp in enumerate(doc.paragraphs):
                if pp._element is p._element:
                    cursor[0] = i
                    break

        insert(f"Step {step_num}. {args.step}", style=SUBHEADING_STYLE)
        if args.how:
            insert("How to do this:", bold=True, size=9, color=RGBColor(0x55, 0x55, 0x55))
            insert(args.how, italic=True, size=9, color=RGBColor(0x3A, 0x3A, 0x3A), left_indent=18)
        if args.result:
            insert("Result:", bold=True, size=9, color=RGBColor(0x55, 0x55, 0x55))
            insert(args.result, size=10, left_indent=18)
        if args.conclusion:
            insert("What this told us:", bold=True, size=9, color=RGBColor(0x55, 0x55, 0x55))
            insert(args.conclusion, size=10, left_indent=18)
        if args.outcome:
            insert(f"Outcome: {args.outcome}", bold=True, size=10, color=RGBColor(0x1F, 0x3A, 0x5F))
        insert("")  # spacer

    append_to_section(doc, "Diagnostic Steps", render)
    save_doc(doc, doc_path)
    print(f"[OK] Added step {step_num}")


def cmd_set_flow(args, doc_path: Path) -> None:
    doc = open_doc(doc_path)
    insert_at = delete_section_body(doc, "How I Worked Through It")
    if insert_at < 0:
        sys.exit("How I Worked Through It heading missing")
    cursor = insert_at - 1
    # Split flow into paragraphs on blank lines, render each as its own paragraph.
    for para in [p.strip() for p in args.flow.split("\n\n") if p.strip()]:
        p = insert_paragraph_after(doc, cursor, para, size=10)
        for i, pp in enumerate(doc.paragraphs):
            if pp._element is p._element:
                cursor = i
                break
    save_doc(doc, doc_path)
    print("[OK] Set flow narrative")


def cmd_add_image(args, doc_path: Path) -> None:
    """Embed an image at the end of a named section, with a caption beneath."""
    from docx.shared import Inches, Pt, RGBColor

    image_path = Path(args.image_path).expanduser().resolve()
    if not image_path.exists():
        sys.exit(f"image not found: {image_path}")

    doc = open_doc(doc_path)

    def render(after_idx: int) -> None:
        cursor = [after_idx]

        # Spacer line so the image isn't flush with the previous content.
        spacer = insert_paragraph_after(doc, cursor[0], "")
        for i, pp in enumerate(doc.paragraphs):
            if pp._element is spacer._element:
                cursor[0] = i
                break

        # Image paragraph.
        img_p = insert_paragraph_after(doc, cursor[0], "")
        try:
            img_p.add_run().add_picture(str(image_path), width=Inches(5))
        except Exception as exc:  # noqa: BLE001
            sys.exit(f"failed to embed image: {exc}")
        for i, pp in enumerate(doc.paragraphs):
            if pp._element is img_p._element:
                cursor[0] = i
                break

        # Caption.
        if args.caption:
            cap = insert_paragraph_after(doc, cursor[0], "")
            run = cap.add_run(f"Figure: {clean(args.caption)}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    append_to_section(doc, args.section, render)
    save_doc(doc, doc_path)
    print(f"[OK] Embedded image in section '{args.section}'")


def cmd_finalize(args, doc_path: Path) -> None:
    doc = open_doc(doc_path)
    for heading, value in [
        ("Root Cause", args.root_cause),
        ("Resolution", args.resolution),
        ("Lesson Learned", args.lesson),
    ]:
        insert_at = delete_section_body(doc, heading)
        if insert_at < 0:
            continue
        insert_paragraph_after(doc, insert_at - 1, value or "(none)", size=10)
    save_doc(doc, doc_path)
    print("[OK] Finalized root cause / resolution / lesson")


# ----------------------- main -----------------------


def main() -> None:
    p = argparse.ArgumentParser()
    p.add_argument("--doc-path", required=True, help="Absolute path to the case .docx")
    sub = p.add_subparsers(dest="action", required=True)

    sp = sub.add_parser("start")
    sp.add_argument("--case-id", required=True)
    sp.add_argument("--client", required=True)
    sp.add_argument("--engineer", default="Visal")
    sp.add_argument("--symptom", default="")
    sp.add_argument("--environment", default="")

    sp = sub.add_parser("add-qa")
    sp.add_argument("--question", required=True)
    sp.add_argument("--answer", default="")

    sp = sub.add_parser("set-causes")
    sp.add_argument("--cause", action="append", default=[], help="A possible cause. Repeat for each.")
    sp.add_argument("--reason", action="append", default=[], help="One-line reasoning for the paired cause. Repeat in same order.")

    sp = sub.add_parser("add-step")
    sp.add_argument("--step", required=True)
    sp.add_argument("--how", default="", help="How a Level 1 should actually run this check")
    sp.add_argument("--result", default="", help="What Visal saw when he ran it")
    sp.add_argument("--conclusion", default="", help="What that result ruled in or out")
    sp.add_argument("--outcome", default="", help="resolved / next step / escalate")

    sp = sub.add_parser("set-flow")
    sp.add_argument("--flow", required=True, help="First-person narrative")

    sp = sub.add_parser("add-image")
    sp.add_argument(
        "--section",
        required=True,
        help="Section heading to embed the image under (e.g. 'Conversation Log', 'Diagnostic Steps')",
    )
    sp.add_argument(
        "--image-path",
        required=True,
        dest="image_path",
        help="Absolute path to the image file",
    )
    sp.add_argument("--caption", default="", help="Short caption shown below the image")

    sp = sub.add_parser("finalize")
    sp.add_argument("--root-cause", default="")
    sp.add_argument("--resolution", default="")
    sp.add_argument("--lesson", default="")

    args = p.parse_args()
    doc_path = Path(args.doc_path).expanduser().resolve()

    dispatch = {
        "start": cmd_start,
        "add-qa": cmd_add_qa,
        "set-causes": cmd_set_causes,
        "add-step": cmd_add_step,
        "set-flow": cmd_set_flow,
        "add-image": cmd_add_image,
        "finalize": cmd_finalize,
    }
    dispatch[args.action](args, doc_path)


if __name__ == "__main__":
    main()
