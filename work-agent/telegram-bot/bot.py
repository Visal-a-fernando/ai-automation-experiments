"""
Telegram bridge for the work-agent.

Each incoming Telegram message is forwarded to `claude -p` running in the
work-agent directory, so the same skills (troubleshooting-wizard,
ms-stack-quickref, learning-planner) that you use locally answer from your
phone. Auth piggybacks on whatever account `claude login` used on this
machine — typically your Max plan.

Run locally (polling), one process, one user (you).
"""

from __future__ import annotations

import asyncio
import base64
import json
import logging
import os
import re
import shutil
import time
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from telegram import Update
from telegram.constants import ChatAction
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    filters,
)

load_dotenv()

BOT_TOKEN = os.environ["TELEGRAM_BOT_TOKEN"]
ALLOWED_USER_IDS = {
    int(uid.strip())
    for uid in os.environ.get("TELEGRAM_ALLOWED_USER_IDS", "").split(",")
    if uid.strip()
}
WORK_AGENT_DIR = Path(
    os.environ.get("WORK_AGENT_DIR", Path(__file__).resolve().parent.parent)
)
CASE_LOG_DIR = WORK_AGENT_DIR / "case-log"
CASE_DOC_DIR = Path(
    os.environ.get("CASE_DOC_DIR", WORK_AGENT_DIR / "case-docs")
)
PDF_WATCH_DIRS = [
    WORK_AGENT_DIR,                  # ad-hoc PDFs (guides, references)
    WORK_AGENT_DIR / "case-log",     # wizard case records
    WORK_AGENT_DIR / "study-plans",  # learning planner outputs
]
# Default to OneDrive so every Telegram photo lands in the synced folder
# alongside the docs that reference it. Path can have spaces — we no longer
# rely on the @path parser, the image goes through stream-json as base64.
INBOX_DIR = Path(os.environ.get("INBOX_DIR", CASE_DOC_DIR / "images"))
INBOX_DIR.mkdir(parents=True, exist_ok=True)
CASE_DOC_DIR.mkdir(exist_ok=True)
CLAUDE_CLI = os.environ.get("CLAUDE_CLI", "claude")
CLAUDE_TIMEOUT_SECONDS = int(os.environ.get("CLAUDE_TIMEOUT_SECONDS", "300"))
TELEGRAM_MAX_CHARS = 4000

logging.basicConfig(
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    level=logging.INFO,
)
log = logging.getLogger("work-agent-bot")

# chat_id -> last claude session_id (resumes the conversation)
SESSIONS: dict[int, str] = {}
# chat_id -> path to the live case doc the wizard is writing into
DOC_PATHS: dict[int, Path] = {}
# chat_id -> path to the flat chat log doc, always growing in chat or case mode
CHAT_DOCS: dict[int, Path] = {}
# chat_id -> "chat" | "case"; defaults to "chat" when key is missing
MODES: dict[int, str] = {}


def _allocate_chat_doc_path() -> Path:
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    return CASE_DOC_DIR / f"CHAT-{ts}.docx"


def _append_chat_turn(
    doc_path: Path,
    user_msg: str,
    assistant_msg: str,
    image_paths: list[Path] | None = None,
) -> None:
    """Append one round of (user, assistant) to the chat log doc.

    If image_paths is non-empty, embed each image inline between the user
    line and the agent line so the doc reads like the Telegram thread.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        log.warning("python-docx not installed, skipping chat log")
        return

    if doc_path.exists():
        doc = Document(str(doc_path))
    else:
        doc = Document()
        title = doc.add_paragraph()
        run = title.add_run("Work Agent Chat Log")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        meta = doc.add_paragraph()
        meta_run = meta.add_run(
            f"Started {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    now = datetime.now().strftime("%H:%M")

    you_p = doc.add_paragraph()
    you_label = you_p.add_run(f"You [{now}]: ")
    you_label.bold = True
    you_label.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    you_p.add_run(user_msg or "(photo only)")

    # Embed each image inline, right after the user line.
    for p in image_paths or []:
        try:
            img_p = doc.add_paragraph()
            img_p.add_run().add_picture(str(p), width=Inches(4))
        except Exception:  # noqa: BLE001
            log.exception("failed to embed image %s in chat doc", p)

    agent_p = doc.add_paragraph()
    agent_label = agent_p.add_run(f"Agent [{now}]: ")
    agent_label.bold = True
    agent_label.font.color.rgb = RGBColor(0x33, 0x66, 0x33)
    agent_p.add_run(assistant_msg)

    doc.add_paragraph()  # blank line spacer

    try:
        doc_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(doc_path))
    except PermissionError:
        log.warning(
            "chat log %s is open in Word on this laptop, skipping this turn",
            doc_path,
        )


def _allocate_doc_path() -> Path:
    """Pick a fresh CASE-YYYY-NNNN.docx filename in CASE_DOC_DIR."""
    year = datetime.now().year
    existing = list(CASE_DOC_DIR.glob(f"CASE-{year}-*.docx"))
    highest = 0
    for f in existing:
        try:
            n = int(f.stem.split("-")[2])
            highest = max(highest, n)
        except (IndexError, ValueError):
            continue
    case_id = f"CASE-{year}-{(highest + 1):04d}"
    return CASE_DOC_DIR / f"{case_id}.docx"


def _doc_footer(doc_path: Path) -> str:
    """Short reference appended after the user message so the wizard knows
    where to write. Kept brief and placed AFTER the user content so the
    model treats the problem description as the primary input."""
    case_id = doc_path.stem
    return (
        f"\n\n---\n"
        f"(case doc: {doc_path.as_posix()} | case id: {case_id})"
    )


def _resolve_claude_cli() -> str:
    """Return an absolute path to the claude executable, or the bare name."""
    found = shutil.which(CLAUDE_CLI)
    return found or CLAUDE_CLI


def _is_authorized(update: Update) -> bool:
    if not ALLOWED_USER_IDS:
        return True  # no whitelist configured — open mode, dev only
    user = update.effective_user
    return user is not None and user.id in ALLOWED_USER_IDS


_IMAGE_MIME = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


def _image_content_block(path: Path) -> dict:
    media_type = _IMAGE_MIME.get(path.suffix.lower(), "image/jpeg")
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return {
        "type": "image",
        "source": {"type": "base64", "media_type": media_type, "data": data},
    }


async def _run_claude_text(prompt: str, session_id: str | None) -> tuple[str, str | None]:
    """Text-only path. Single-shot --output-format json, no streaming needed."""
    cmd = [
        _resolve_claude_cli(),
        "-p",
        "--output-format", "json",
        "--permission-mode", "bypassPermissions",
    ]
    if session_id:
        cmd += ["--resume", session_id]
    cmd.append(prompt)

    proc = await asyncio.create_subprocess_exec(
        *cmd,
        cwd=str(WORK_AGENT_DIR),
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    try:
        stdout, stderr = await asyncio.wait_for(
            proc.communicate(), timeout=CLAUDE_TIMEOUT_SECONDS
        )
    except asyncio.TimeoutError:
        proc.kill()
        await proc.wait()
        raise RuntimeError(f"claude did not respond within {CLAUDE_TIMEOUT_SECONDS}s")

    if proc.returncode != 0:
        err = stderr.decode("utf-8", errors="replace").strip()
        raise RuntimeError(f"claude exited {proc.returncode}: {err[:500]}")

    raw = stdout.decode("utf-8", errors="replace").strip()
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError:
        return raw, session_id
    text = payload.get("result") or payload.get("response") or raw
    new_session = payload.get("session_id") or session_id
    return text.strip(), new_session


async def _run_claude_stream(
    prompt: str,
    session_id: str | None,
    image_paths: list[Path],
) -> tuple[str, str | None]:
    """Stream-json path used when there are images to attach. Reads stdout
    line-by-line so the OS pipe never backs up, which was the root cause of
    the earlier 5-minute hang."""
    cmd = [
        _resolve_claude_cli(),
        "-p",
        "--input-format", "stream-json",
        "--output-format", "stream-json",
        "--verbose",
        "--permission-mode", "bypassPermissions",
    ]
    if session_id:
        cmd += ["--resume", session_id]

    # Tell the wizard where the image lives on disk so it can embed it in
    # the case doc via update_case_doc.py add-image. The image is also passed
    # as a base64 content block below so the model can SEE it.
    path_lines = "\n".join(f"  - {p.as_posix()}" for p in image_paths)
    prompt_with_paths = (
        f"{prompt}\n\n"
        f"[Image file(s) attached, saved at:\n{path_lines}\n"
        f"If you are in case mode, after acknowledging what you see, embed each "
        f"image in the case doc with: update_case_doc.py add-image "
        f"--doc-path <case doc> --section <section name> --image-path <path> "
        f"--caption \"<one line describing what it shows>\". Use "
        f"\"Conversation Log\" for intake photos, \"Diagnostic Steps\" for "
        f"result screenshots.]"
    )

    content_blocks: list[dict] = [{"type": "text", "text": prompt_with_paths}]
    for p in image_paths:
        try:
            content_blocks.append(_image_content_block(p))
        except Exception:  # noqa: BLE001
            log.exception("failed to encode image %s, skipping", p)
    msg = {"type": "user", "message": {"role": "user", "content": content_blocks}}
    stdin_bytes = (json.dumps(msg) + "\n").encode("utf-8")

    proc = await asyncio.create_subprocess_exec(
        *cmd,
        cwd=str(WORK_AGENT_DIR),
        stdin=asyncio.subprocess.PIPE,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )

    assert proc.stdin is not None and proc.stdout is not None and proc.stderr is not None
    proc.stdin.write(stdin_bytes)
    await proc.stdin.drain()
    proc.stdin.close()

    text = ""
    new_session = session_id
    stderr_buf = bytearray()

    async def consume_stdout() -> None:
        nonlocal text, new_session
        # asyncio.StreamReader.readline yields one line at a time.
        while True:
            line = await proc.stdout.readline()
            if not line:
                break
            try:
                event = json.loads(line.decode("utf-8", errors="replace"))
            except json.JSONDecodeError:
                continue
            etype = event.get("type")
            if etype == "result":
                text = event.get("result") or text
                new_session = event.get("session_id") or new_session
            elif etype == "system" and event.get("session_id"):
                new_session = event.get("session_id")

    async def consume_stderr() -> None:
        while True:
            chunk = await proc.stderr.read(4096)
            if not chunk:
                break
            stderr_buf.extend(chunk)

    try:
        await asyncio.wait_for(
            asyncio.gather(consume_stdout(), consume_stderr()),
            timeout=CLAUDE_TIMEOUT_SECONDS,
        )
    except asyncio.TimeoutError:
        proc.kill()
        await proc.wait()
        raise RuntimeError(f"claude did not respond within {CLAUDE_TIMEOUT_SECONDS}s")

    await proc.wait()
    if proc.returncode != 0:
        err = stderr_buf.decode("utf-8", errors="replace").strip()
        raise RuntimeError(f"claude exited {proc.returncode}: {err[:500]}")

    return (text or "(empty response)").strip(), new_session


async def _run_claude(
    prompt: str,
    session_id: str | None,
    image_paths: list[Path] | None = None,
) -> tuple[str, str | None]:
    """Route text-only vs image turns to the right runner."""
    log.info(
        "running claude (resume=%s, prompt_len=%d, images=%d)",
        session_id,
        len(prompt),
        len(image_paths or []),
    )
    if image_paths:
        return await _run_claude_stream(prompt, session_id, image_paths)
    return await _run_claude_text(prompt, session_id)


_MD_BOLD = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)
_MD_UNDERLINE = re.compile(r"__(.+?)__", re.DOTALL)
_MD_HEADER = re.compile(r"^[ \t]*#{1,6}[ \t]+", re.MULTILINE)
# Single * around text. Careful to require non-space, non-* on inside edges so
# we don't strip glob characters in cmdlets or list bullets at line start.
_MD_ITALIC_STAR = re.compile(r"(?<![\w*])\*(?!\s)([^\*\n]+?)(?<!\s)\*(?![\w*])")
# Single _ around text. Skip identifiers like step_1_result by requiring word
# boundary anchors outside.
_MD_ITALIC_UNDER = re.compile(r"(?<![\w_])_(?!\s)([^_\n]+?)(?<!\s)_(?![\w_])")


def _strip_markdown(text: str) -> str:
    """Remove markdown markers that Telegram renders as literal characters."""
    text = _MD_BOLD.sub(r"\1", text)
    text = _MD_UNDERLINE.sub(r"\1", text)
    text = _MD_ITALIC_STAR.sub(r"\1", text)
    text = _MD_ITALIC_UNDER.sub(r"\1", text)
    text = _MD_HEADER.sub("", text)
    return text


def _strip_dashes(text: str) -> str:
    """Visal's house rule: no em / en / double-hyphen dashes anywhere."""
    return (
        text.replace(" — ", ", ")
        .replace(" – ", ", ")
        .replace("—", ", ")
        .replace("–", " to ")
        .replace(" -- ", ", ")
        .replace("--", ", ")
    )


async def _send_long(update: Update, text: str, already_cleaned: bool = False) -> None:
    if not text:
        text = "(empty response)"
    if not already_cleaned:
        text = _strip_dashes(_strip_markdown(text))
    for i in range(0, len(text), TELEGRAM_MAX_CHARS):
        chunk = text[i : i + TELEGRAM_MAX_CHARS]
        await update.effective_message.reply_text(chunk)


def _snapshot_pdfs() -> set[Path]:
    found: set[Path] = set()
    for d in PDF_WATCH_DIRS:
        if d.exists():
            found.update(p.resolve() for p in d.glob("*.pdf"))
    return found


async def _send_new_pdfs(update: Update, before: set[Path]) -> None:
    after = _snapshot_pdfs()
    new_pdfs = sorted(after - before, key=lambda p: p.stat().st_mtime)
    for pdf in new_pdfs:
        try:
            rel = pdf.relative_to(WORK_AGENT_DIR)
        except ValueError:
            rel = pdf
        try:
            with open(pdf, "rb") as fh:
                await update.effective_message.reply_document(
                    document=fh,
                    filename=pdf.name,
                    caption=f"Saved to {rel}",
                )
        except Exception:  # noqa: BLE001
            log.exception("failed to send pdf %s", pdf)


async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    if not _is_authorized(update):
        await update.effective_message.reply_text(
            f"Not authorized. Your Telegram ID is {user.id if user else '?'} — "
            "add it to TELEGRAM_ALLOWED_USER_IDS to enable."
        )
        return
    await update.effective_message.reply_text(
        "Work-agent bridge online. Defaults to chat mode, every turn is logged "
        "to a Word doc in OneDrive.\n\n"
        "/case  start a structured ticket triage (separate case doc)\n"
        "/chat  switch back to chat mode\n"
        "/new   fresh session, new chat log\n"
        "/whoami show your Telegram user ID"
    )


async def cmd_new(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _is_authorized(update):
        return
    chat_id = update.effective_chat.id
    SESSIONS.pop(chat_id, None)
    DOC_PATHS.pop(chat_id, None)
    CHAT_DOCS.pop(chat_id, None)
    MODES[chat_id] = "chat"
    await update.effective_message.reply_text(
        "Fresh start. Chat mode, new chat log will be created on your next "
        "message. Send /case when you want to open a structured ticket."
    )


async def cmd_case(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _is_authorized(update):
        return
    chat_id = update.effective_chat.id
    MODES[chat_id] = "case"
    SESSIONS.pop(chat_id, None)
    DOC_PATHS.pop(chat_id, None)
    await update.effective_message.reply_text(
        "Case mode on. Drop the ticket details and the wizard will start a "
        "fresh case doc and walk you through it. /chat to switch back."
    )


async def cmd_chat(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _is_authorized(update):
        return
    chat_id = update.effective_chat.id
    MODES[chat_id] = "chat"
    DOC_PATHS.pop(chat_id, None)
    await update.effective_message.reply_text(
        "Chat mode on. Ask me anything, every turn lands in your chat log. "
        "/case when you want a structured ticket."
    )


async def cmd_whoami(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    await update.effective_message.reply_text(
        f"Telegram user ID: {user.id if user else '?'}"
    )


async def _download_telegram_image(
    ctx: ContextTypes.DEFAULT_TYPE, update: Update
) -> Path:
    """Download whichever image attachment is on the message (photo or image doc)."""
    msg = update.effective_message
    if msg.photo:
        attachment = msg.photo[-1]  # largest size
        ext = "jpg"
    elif msg.document and (msg.document.mime_type or "").startswith("image/"):
        attachment = msg.document
        # Reuse the original extension if we can, else fall back to jpg.
        original = (msg.document.file_name or "").rsplit(".", 1)
        ext = original[-1] if len(original) == 2 and len(original[-1]) <= 5 else "jpg"
    else:
        raise RuntimeError("no image attachment on this message")

    tg_file = await ctx.bot.get_file(attachment.file_id)
    fname = f"img-{int(time.time() * 1000)}-{attachment.file_unique_id}.{ext}"
    dest = (INBOX_DIR / fname).resolve()
    await tg_file.download_to_drive(custom_path=str(dest))
    return dest


async def on_photo(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _is_authorized(update):
        await update.effective_message.reply_text("Not authorized.")
        return

    chat_id = update.effective_chat.id
    caption = (update.effective_message.caption or "").strip()

    await ctx.bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
    try:
        photo_path = await _download_telegram_image(ctx, update)
    except Exception as exc:  # noqa: BLE001
        log.exception("photo download failed")
        await update.effective_message.reply_text(f"Couldn't download the photo: {exc}")
        return

    log.info("image downloaded: %s (%d bytes)", photo_path, photo_path.stat().st_size)
    intro = caption or "Photo attached from my phone. Read it and continue troubleshooting."
    await _dispatch_prompt(update, ctx, intro, image_paths=[photo_path])


async def on_message(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _is_authorized(update):
        await update.effective_message.reply_text("Not authorized.")
        return

    prompt = (update.effective_message.text or "").strip()
    if not prompt:
        return

    await ctx.bot.send_chat_action(
        chat_id=update.effective_chat.id, action=ChatAction.TYPING
    )
    await _dispatch_prompt(update, ctx, prompt)


async def _dispatch_prompt(
    update: Update,
    ctx: ContextTypes.DEFAULT_TYPE,
    prompt: str,
    image_paths: list[Path] | None = None,
) -> None:
    chat_id = update.effective_chat.id
    pdfs_before = _snapshot_pdfs()
    mode = MODES.get(chat_id, "chat")
    user_msg = prompt  # save before we wrap it with preambles

    # Always make sure there's a chat log doc allocated for this session.
    if chat_id not in CHAT_DOCS:
        CHAT_DOCS[chat_id] = _allocate_chat_doc_path()
        log.info("allocated chat log %s for chat %s", CHAT_DOCS[chat_id], chat_id)

    if mode == "case":
        if chat_id not in DOC_PATHS:
            DOC_PATHS[chat_id] = _allocate_doc_path()
            log.info(
                "allocated case doc %s for chat %s", DOC_PATHS[chat_id], chat_id
            )
        prompt = prompt + _doc_footer(DOC_PATHS[chat_id])

    try:
        text, new_session = await _run_claude(
            prompt, SESSIONS.get(chat_id), image_paths=image_paths
        )
    except FileNotFoundError:
        await update.effective_message.reply_text(
            "claude CLI not found on PATH. Install it with "
            "`npm install -g @anthropic-ai/claude-code` and run `claude login`."
        )
        return
    except Exception as exc:  # noqa: BLE001
        log.exception("claude invocation failed")
        await update.effective_message.reply_text(f"Error: {exc}")
        return

    if new_session:
        SESSIONS[chat_id] = new_session

    cleaned = _strip_markdown(text)
    cleaned = _strip_dashes(cleaned)

    # Log the turn to the chat doc before sending, so a failed Telegram send
    # still gets captured in OneDrive.
    try:
        _append_chat_turn(
            CHAT_DOCS[chat_id], user_msg, cleaned, image_paths=image_paths
        )
    except Exception:  # noqa: BLE001
        log.exception("failed to append to chat doc")

    await _send_long(update, cleaned, already_cleaned=True)
    await _send_new_pdfs(update, pdfs_before)


def main() -> None:
    if not WORK_AGENT_DIR.exists():
        raise SystemExit(f"WORK_AGENT_DIR does not exist: {WORK_AGENT_DIR}")
    log.info("work-agent dir: %s", WORK_AGENT_DIR)
    log.info("allowed user ids: %s", ALLOWED_USER_IDS or "(open)")

    app = ApplicationBuilder().token(BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("new", cmd_new))
    app.add_handler(CommandHandler("case", cmd_case))
    app.add_handler(CommandHandler("chat", cmd_chat))
    app.add_handler(CommandHandler("whoami", cmd_whoami))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, on_message))
    app.add_handler(MessageHandler(filters.PHOTO, on_photo))
    app.add_handler(MessageHandler(filters.Document.IMAGE, on_photo))
    log.info("bot polling…")
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
